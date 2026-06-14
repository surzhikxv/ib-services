# -*- coding: utf-8 -*-
"""Сухой чек-лист Phase 1, разделённый на: нужно с вас (клиент) / нужно с меня (разработчик)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.12

for sec in doc.sections:
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)


def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLACK


def part(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = BLACK
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def head(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = BLACK


def b(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


title("Что нужно для Phase 1")

# ============================ С ВАС ============================
part("Нужно с вас")

head("Доступы к площадкам")
b("YouTube — Google-аккаунт и OAuth-доступ к каналу.")
b("VK — доступ к сообществу/аккаунту, токен.")
b("Telegram — токен бота и права администратора канала.")
b("TikTok — доступ к аккаунту.")
b("Instagram — доступ к аккаунту.")

head("BotHelp")
b("API-ключ.")
b("Тариф с доступом к API (проверить, при необходимости апгрейд).")
b("Доступ к кабинету.")

head("TikTok и Instagram (данные)")
b("Тип аккаунтов: Business/Creator или обычные.")
b("Образец выгрузки статистики.")
b("Регулярная подгрузка выгрузок после запуска.")

head("Оплата сервисов")
b("Хостинг (VPS) и домен.")
b("Аккаунт и ключ языковой модели.")
b("Ежемесячные расходы на хостинг и API.")

head("Бизнес-данные и решения")
b("Структура тарифов курса.")
b("Список конкурентов (5–10).")
b("Критерии подбора партнёров для коллабов.")
b("Какие метрики в приоритете и что считать успехом.")
b("Тон и формат отчётов ИИ.")

head("Участие по ходу")
b("Показать структуру воронки и тегов в BotHelp.")
b("Оценить ИИ-разборы на своих данных и дать правки.")
b("Решение по обходу блокировок при парсинге: прокси или ручной ввод.")
b("Приёмка результата.")

head("Постоянно после запуска")
b("Загрузка выгрузок TikTok и Instagram.")
b("Финальный выбор и отправка сообщений партнёрам.")

# ============================ С МЕНЯ ============================
part("Нужно с меня")

head("Озеро данных")
b("Коннекторы YouTube, VK, Telegram.")
b("Парсеры выгрузок TikTok и Instagram.")
b("Интеграция с BotHelp: подписчики, теги, шаги воронки, оплаты Prodamus.")
b("Единая база и нормализация данных.")

head("Воронка")
b("Маппинг тегов и шагов BotHelp в этапы воронки.")
b("Атрибуция источника, настройка старт-ссылок и меток.")

head("Дашборд")
b("Metabase поверх базы: панели по воронке, выручке, динамике.")

head("ИИ-аналитик")
b("Пайплайн данных и промпты.")
b("Недельный разбор и режим вопросов по своим данным.")
b("Доставка в Telegram-бот и дашборд.")

head("Коллабы и ниша")
b("Подбор партнёров и черновики сообщений.")
b("Мониторинг конкурентов и ИИ-разбор.")

head("Запуск и поддержка")
b("Сценарии n8n и расписания.")
b("Развёртывание на сервере и резервное копирование.")
b("Передача доступов и инструкции.")
b("Устранение багов в гарантийный период.")

out = "/Users/seva/Desktop/инфобизнес/Что нужно — с вас и с меня (Phase 1).docx"
doc.save(out)
print("Saved:", out)
